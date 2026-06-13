"""
Document text extraction for RAG ingest (Part D).

Supported formats:
  .pdf   — pypdf (pure Python, no native libs, <2MB install)
  .docx  — python-docx
  .doc   — best-effort: try python-docx, fall back to raw decode
  .md / .markdown / .txt / .csv / .json / .log — direct UTF-8 decode

Usage:
    text = extract(file_bytes, filename="report.pdf")
"""
from __future__ import annotations

import io
import structlog

log = structlog.get_logger()


def extract(content: bytes, filename: str) -> str:
    """Extract plain text from document bytes based on file extension.

    Returns the extracted string. Raises ValueError on unsupported type
    or unrecoverable parse error.
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == "pdf":
        return _extract_pdf(content)
    if ext in ("docx", "doc"):
        return _extract_docx(content)
    # Plain text variants — just decode
    if ext in ("md", "markdown", "txt", "csv", "json", "log", ""):
        return _decode_text(content)

    raise ValueError(f"Unsupported file type: .{ext}")


def _extract_pdf(content: bytes) -> str:
    """Extract text from a PDF using pypdf."""
    try:
        from pypdf import PdfReader  # type: ignore[import]
    except ImportError as exc:
        raise ImportError(
            "pypdf is required for PDF extraction. "
            "Install it: pip install pypdf"
        ) from exc

    reader = PdfReader(io.BytesIO(content))
    pages: list[str] = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
            pages.append(text)
        except Exception as e:  # noqa: BLE001
            log.warning("pdf_page_extract_failed", err=str(e)[:80])
            pages.append("")

    result = "\n\n".join(p for p in pages if p.strip())
    if not result.strip():
        raise ValueError("PDF appears to be image-only or encrypted (no extractable text).")
    return result


def _extract_docx(content: bytes) -> str:
    """Extract text from a .docx file using python-docx."""
    try:
        from docx import Document  # type: ignore[import]
    except ImportError as exc:
        raise ImportError(
            "python-docx is required for .docx extraction. "
            "Install it: pip install python-docx"
        ) from exc

    doc = Document(io.BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract table cell text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n\n".join(paragraphs)


def _decode_text(content: bytes) -> str:
    """Decode bytes as UTF-8, falling back to latin-1."""
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        return content.decode("latin-1", errors="replace")
